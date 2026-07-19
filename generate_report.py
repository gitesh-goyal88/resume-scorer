"""
generate_report.py — ResumeIQ B.Tech Summer Internship Report Generator
Generates a 14-page JIIT-format Word document.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

IMG_DIR = "assets/images_png"
OUT_PATH = os.path.expanduser("~/Desktop/ResumeIQ_Internship_Report.docx")

doc = Document()

# ── Global Page Setup ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── Style helpers ──────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, center=False, size=14, bold=True, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    return p

def body(text, justify=True, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def bullet(text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=size)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_img(path, width=Inches(5.5)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        p.paragraph_format.space_after = Pt(8)

def divider():
    p = doc.add_paragraph("─" * 80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(180, 180, 180)

def section_header(text):
    """Bold underlined section heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.underline = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)

def code_block(lines):
    """Monospace code snippet block"""
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        set_font(run, name="Courier New", size=9)
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("JAYPEE INSTITUTE OF INFORMATION TECHNOLOGY")
set_font(run, size=16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("(Deemed to be University under Section 3 of UGC Act 1956)")
set_font(run, size=11, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sector 62, Noida – 201309, Uttar Pradesh, India")
set_font(run, size=11)

doc.add_paragraph()
divider()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SUMMER INTERNSHIP REPORT")
set_font(run, size=15, bold=True)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ResumeIQ: An AI-Powered ATS Resume Optimizer,\nLive Resume Builder & Intelligent Job Matcher")
set_font(run, size=14, bold=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted By")
set_font(run, size=12, bold=True)

table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
details = [
    ("Gitesh Goyal", "22103055"),
    ("Arjav Jain",   "22103XXX"),
    ("Riddhi",       "22103XXX"),
]
for i, (name, roll) in enumerate(details):
    row = table.rows[i]
    for cell, txt in zip(row.cells, [name, roll]):
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(txt)
        set_font(run2, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("B.Tech — Computer Science and Engineering (AI & ML)\nBatch: 2022–2026")
set_font(run, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Guided By:\nDr. [Faculty Mentor Name]\nDepartment of Computer Science & Information Technology\nJIIT, Noida")
set_font(run, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("June – July 2026")
set_font(run, size=12, bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — DECLARATION
# ══════════════════════════════════════════════════════════════════════════════
heading("DECLARATION", center=True, size=14)
divider()
doc.add_paragraph()

body(
    'We hereby declare that the Summer Internship Project titled ResumeIQ: An AI-Powered '
    'ATS Resume Optimizer, Live Resume Builder & Intelligent Job Matcher, submitted to the '
    'Department of Computer Science & Information Technology, JIIT Noida, is a record of '
    'original work carried out by us during the Summer Internship period of June-July 2026.'
)
body(
    'The work presented in this report has not been submitted elsewhere for the award of any '
    'degree or diploma. All information and results are based on our own research and '
    'implementation. Wherever the work of others has been used, it has been duly acknowledged.'
)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

table = doc.add_table(rows=2, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
names = ["Gitesh Goyal\n22103055", "Arjav Jain\n22103XXX", "Riddhi\n22103XXX"]
for i, name in enumerate(names):
    cell = table.rows[0].cells[i]
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("__________________\n" + name)
    set_font(run2, size=11)

cell2 = table.rows[1].cells[0]
p3 = cell2.paragraphs[0]
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
run3 = p3.add_run("Date: _______________")
set_font(run3, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — CERTIFICATE
# ══════════════════════════════════════════════════════════════════════════════
heading("CERTIFICATE", center=True, size=14)
divider()
doc.add_paragraph()

body(
    'This is to certify that the Summer Internship Project titled ResumeIQ: An AI-Powered '
    'ATS Resume Optimizer, Live Resume Builder & Intelligent Job Matcher has been successfully '
    'completed by the following students of B.Tech (Computer Science & Engineering - AI & ML), '
    'Batch 2022-2026, JIIT Noida, as part of their Summer Internship Programme (June-July 2026).'
)

doc.add_paragraph()
bullet("Gitesh Goyal — Enrollment No.: 22103055")
bullet("Arjav Jain   — Enrollment No.: 22103XXX")
bullet("Riddhi       — Enrollment No.: 22103XXX")
doc.add_paragraph()

body(
    "The project has been carried out under my supervision and guidance. The work "
    "reported in this document is original and has not been submitted elsewhere for "
    "the award of any degree or diploma."
)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Dr. [Faculty Mentor Name]")
set_font(run, size=12, bold=True)

p = doc.add_paragraph()
run = p.add_run(
    "Assistant Professor / Associate Professor\n"
    "Department of Computer Science & Information Technology\n"
    "Jaypee Institute of Information Technology, Noida"
)
set_font(run, size=12)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Date: _______________         Place: Noida")
set_font(run, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading("TABLE OF CONTENTS", center=True, size=14)
divider()
doc.add_paragraph()

toc = [
    ("1.",  "Abstract",                          "5"),
    ("2.",  "Introduction",                      "6"),
    ("3.",  "Problem Statement & Objectives",    "7"),
    ("4.",  "Methodology",                       "8"),
    ("5.",  "Implementation & Code Snippets",    "9"),
    ("6.",  "Results & Discussion",              "10–11"),
    ("7.",  "Learning Outcomes",                 "12"),
    ("8.",  "Conclusion",                        "13"),
    ("9.",  "Future Scope",                      "13"),
    ("10.", "References",                        "14"),
]

tbl = doc.add_table(rows=len(toc), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (num, title, page) in enumerate(toc):
    row = tbl.rows[i]
    for j, txt in enumerate([num, title, f"Page {page}"]):
        p2 = row.cells[j].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.RIGHT
        run2 = p2.add_run(txt)
        set_font(run2, size=12, bold=(i == 0))
    row.cells[0].width = Inches(0.5)
    row.cells[2].width = Inches(1.2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
heading("1. ABSTRACT", size=14)
divider()
doc.add_paragraph()

body(
    "ResumeIQ is a full-stack, AI-driven web application built to solve the fundamental "
    "challenge faced by modern job seekers: understanding and improving their resume's "
    "compatibility with Applicant Tracking Systems (ATS). The core intelligence of the "
    "platform is powered by a custom ensemble of scikit-learn Machine Learning models, "
    "including a Multinomial Naive Bayes classifier for job-role prediction, a TF-IDF "
    "Vectorizer with Cosine Similarity for semantic job-to-resume matching, and a "
    "Sentence-Transformer-based dense vector engine for quantitative bullet point evaluation."
)
body(
    "The system processes raw PDF resumes, extracts structured text via PyMuPDF, and feeds "
    "this through the trained ML pipeline to produce a comprehensive ATS report card — "
    "complete with a weighted score, skill gap analysis against a live job corpus of 2,800+ "
    "real job descriptions, and granular, line-by-line feedback on every resume bullet point."
)
body(
    "The platform further includes an interactive Live Resume Editor that compiles a "
    "high-fidelity, print-ready PDF in real time, a Smart Job Matcher (using K-Nearest "
    "Neighbors), a JD Matcher (Cosine Similarity on user-pasted job descriptions), an AI "
    "Interview Preparation module, and a global Analytics Dashboard. ResumeIQ is deployed "
    "as an open-source Streamlit web application, backed by a SQLite database of 2,800+ "
    "curated job listings."
)

p = doc.add_paragraph()
run = p.add_run("Keywords: ")
set_font(run, size=12, bold=True)
run2 = p.add_run(
    "ATS Optimization, TF-IDF, Cosine Similarity, Naive Bayes, KNN, "
    "Sentence Transformers, Resume Parsing, Streamlit, Information Retrieval."
)
set_font(run2, size=12, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 6 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading("2. INTRODUCTION", size=14)
divider()
doc.add_paragraph()

section_header("2.1 Background")
body(
    "The modern hiring landscape is increasingly automated. Over 98% of Fortune 500 companies "
    "rely on Applicant Tracking Systems (ATS) to filter inbound applications before any human "
    "recruiter ever reads them. These systems parse resumes and score them against keyword-based "
    "heuristics derived from the job description. A resume with excellent content but poor ATS "
    "compatibility is invisible to recruiters. This creates a critical but invisible barrier for "
    "candidates who lack the technical knowledge to optimize their application documents."
)

section_header("2.2 Motivation")
body(
    "Existing tools for resume optimization are either prohibitively expensive (Jobscan, Rezi), "
    "provide only surface-level keyword matching, or fail to offer actionable, quantifiable "
    "feedback on the quality of individual resume bullet points. There is a clear gap for an "
    "open-source, ML-powered platform that provides deep, technical insight into resume quality "
    "while also serving as a complete job-searching and interview-preparation hub."
)

section_header("2.3 Scope of the Project")
body(
    "ResumeIQ was designed and implemented as a complete, production-ready web application during "
    "the Summer Internship period. The project encompasses: data collection and preprocessing, "
    "training and evaluation of multiple scikit-learn models, building a multi-page interactive "
    "Streamlit UI with a custom glassmorphic CSS design system, integrating a real-time PDF "
    "compilation engine, and deploying the final product to the Streamlit Community Cloud."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 7 — PROBLEM STATEMENT & OBJECTIVES
# ══════════════════════════════════════════════════════════════════════════════
heading("3. PROBLEM STATEMENT & OBJECTIVES", size=14)
divider()
doc.add_paragraph()

section_header("3.1 Problem Statement")
body(
    "Job seekers consistently fail to pass automated ATS screening filters, not due to a lack "
    "of qualifications, but due to poor resume formatting, missing keywords, and weak bullet "
    "point phrasing. There is no free, open-source tool that (a) quantitatively scores resume "
    "quality using trained ML models, (b) performs semantic job matching against a live corpus "
    "of real job listings, and (c) provides actionable, bullet-level feedback using NLP."
)

section_header("3.2 Objectives")
objectives = [
    "Train a Multinomial Naive Bayes classifier on a labeled resume dataset to predict a candidate's job category with >90% accuracy.",
    "Implement a TF-IDF + Cosine Similarity pipeline to match a candidate's resume against 2,800+ real job descriptions and produce a ranked list of semantic matches.",
    "Develop a weighted ATS Health Score engine that evaluates skills, action verbs, quantifiable metrics, section completeness, and formatting quality.",
    "Build a Sentence Transformer + Cosine Similarity engine to provide granular, line-by-line feedback on resume bullet points against a gold-standard dataset.",
    "Create an interactive, real-time PDF Resume Builder where candidates can edit their resume in a structured form and instantly preview the compiled ATS-friendly PDF.",
    "Implement a JD Matcher tab where candidates can paste any job description and receive an instant TF-IDF Cosine Similarity match score along with identified matching and missing keywords.",
    "Build a global Analytics Dashboard that displays model performance metrics (Precision, Recall, F1-Score), a Resume Alignment Heatmap, and a market-skill Radar Chart.",
    "Deploy the full application as an open-source, publicly accessible web application on the Streamlit Community Cloud.",
]
for i, obj in enumerate(objectives, 1):
    bullet(f"O{i}: {obj}")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 8 — METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
heading("4. METHODOLOGY", size=14)
divider()
doc.add_paragraph()

section_header("4.1 Dataset")
body(
    "The primary training dataset is the 'Updated Resume Dataset' (2,484 annotated resumes "
    "across 25 job categories), sourced from Kaggle. This was augmented with a custom-generated "
    "corpus of 2,800+ real job descriptions scraped and curated into a SQLite database. For the "
    "Bullet Point Evaluator, a hand-crafted Gold Standard dataset of 50+ expert-written, "
    "XYZ-formula resume bullets was constructed to serve as the target embedding space."
)

section_header("4.2 ML Pipeline Architecture")
body(
    "The system employs three distinct ML models working in concert, each solving a separate "
    "sub-problem in the resume analysis workflow:"
)

bullet("Model 1 — Job Role Classifier (Multinomial Naive Bayes + TF-IDF): "
       "Resume text is cleaned, vectorized using TF-IDF (max 1500 features, English stopwords "
       "removed), and classified into one of 25 job categories using a MultinomialNB model "
       "trained on a 60-40 train-test split.")
bullet("Model 2 — ATS Health Score Engine (Rule-Based Weighted Scoring): "
       "A deterministic scoring function that evaluates five dimensions — skill count, action "
       "verb count, quantifiable metrics count, section completeness, and formatting penalty — "
       "using weighted aggregation to produce a final score (0–100) and a letter grade.")
bullet("Model 3 — Bullet Point Evaluator (SentenceTransformers + Cosine Similarity): "
       "Candidate bullet points are encoded into 384-dimensional dense vectors using the "
       "all-MiniLM-L6-v2 SentenceTransformer model. These are compared via Cosine Similarity "
       "against a pre-computed Gold Standard matrix to generate an impact score and specific "
       "NLP-driven feedback.")

section_header("4.3 Job Matching Pipeline")
body(
    "The Job Matcher builds a combined TF-IDF corpus from all 2,800+ job descriptions plus "
    "the candidate's resume. A single TF-IDF matrix is fit on this corpus (max 5000 features, "
    "sublinear TF scaling). The resume vector is then compared against all job vectors using "
    "Cosine Similarity. A baseline-relative scaling function normalises raw similarity scores "
    "into user-friendly percentages."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 9 — IMPLEMENTATION & CODE SNIPPETS
# ══════════════════════════════════════════════════════════════════════════════
heading("5. IMPLEMENTATION & CODE SNIPPETS", size=14)
divider()
doc.add_paragraph()

section_header("Snippet 1 — TF-IDF Vectorisation & Naive Bayes Training (ml_model.py)")
body("The following snippet shows the core model training logic for the Job Role Classifier. "
     "Resume text is first cleaned, then vectorized using TF-IDF, and finally fitted to a "
     "Multinomial Naive Bayes classifier on a stratified 80-20 split.")
code_block([
    "# ml_model.py — Model 1: Job Role Classifier",
    "tfidf = TfidfVectorizer(max_features=1500, stop_words='english')",
    "X = tfidf.fit_transform(df['cleaned_resume'])",
    "y = df['Category'].tolist()",
    "",
    "X_train, X_test, y_train, y_test = train_test_split(",
    "    X, y, test_size=0.2, random_state=42",
    ")",
    "",
    "nb_clf = MultinomialNB()",
    "nb_clf.fit(X_train, y_train)",
    "",
    "train_acc = accuracy_score(y_train, nb_clf.predict(X_train))",
    "test_acc  = accuracy_score(y_test,  nb_clf.predict(X_test))",
    "# Results: Train Acc = 95.32% | Test Acc = 91.69%",
])

doc.add_paragraph()
section_header("Snippet 2 — Cosine Similarity Inference for Bullet Evaluation (ml_model.py)")
body("The inference function encodes candidate bullet points using SentenceTransformers "
     "and evaluates them against the pre-computed Gold Standard matrix via Cosine Similarity.")
code_block([
    "# ml_model.py — Model 3: Semantic Bullet Evaluator",
    "from sentence_transformers import SentenceTransformer",
    "from sklearn.metrics.pairwise import cosine_similarity",
    "",
    "model = SentenceTransformer('all-MiniLM-L6-v2')",
    "candidate_matrix = model.encode(candidate_bullets, convert_to_numpy=True)",
    "",
    "# Compare candidates against gold standard (50 expert bullets)",
    "similarity_scores = cosine_similarity(candidate_matrix, gold_matrix)",
    "",
    "for i, bullet in enumerate(candidate_bullets):",
    "    max_sim = float(np.max(similarity_scores[i]))",
    "    normalized_score = min(max_sim / 0.70 * 100, 100)",
    "    label = 'Strong' if normalized_score >= 65 else 'Weak'",
])

doc.add_paragraph()
section_header("Snippet 3 — TF-IDF Job Matching Pipeline (job_matcher.py)")
body("The Job Matcher fits a single TF-IDF matrix across all job descriptions and the "
     "candidate resume, then ranks jobs by cosine similarity using a baseline-relative scaling.")
code_block([
    "# job_matcher.py — TF-IDF + Cosine Similarity Job Matcher",
    "corpus = job_texts + [resume_text]",
    "processed_corpus = [preprocess(text) for text in corpus]",
    "",
    "vectorizer = TfidfVectorizer(max_features=5000, sublinear_tf=True)",
    "tfidf_matrix = vectorizer.fit_transform(processed_corpus)",
    "",
    "resume_vector = tfidf_matrix[-1]",
    "job_vectors   = tfidf_matrix[:-1]",
    "similarities  = cosine_similarity(resume_vector, job_vectors).flatten()",
    "",
    "# Baseline-relative score scaling",
    "max_raw  = similarities.max()",
    "baseline = max(max_raw, 0.30)",
    "ranked_indices = np.argsort(similarities)[::-1]",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 10 — RESULTS & DISCUSSION (Part 1)
# ══════════════════════════════════════════════════════════════════════════════
heading("6. RESULTS & DISCUSSION", size=14)
divider()
doc.add_paragraph()

section_header("6.1 Model Performance — Job Role Classifier")
body(
    "The Naive Bayes classifier (60-40 train-test split) achieves the following metrics on "
    "the 2,484-resume labeled dataset across 25 industry categories:"
)

metrics = [
    ("Train Accuracy",  "95.32%"),
    ("Test Accuracy",   "91.69%"),
    ("Precision",       "96.15%"),
    ("Recall",          "89.78%"),
    ("F1-Score",        "90.74%"),
]
tbl = doc.add_table(rows=len(metrics) + 1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
hdr_cells = tbl.rows[0].cells
for cell, txt in zip(hdr_cells, ["Metric", "Value"]):
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(txt)
    set_font(run2, size=11, bold=True)

for i, (metric, value) in enumerate(metrics):
    row = tbl.rows[i + 1]
    for cell, txt in zip(row.cells, [metric, value]):
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(txt)
        set_font(run2, size=11)

doc.add_paragraph()
body("Figure 6.1 below shows the ML Metrics panel from the live Analytics Dashboard:")
add_img(f"{IMG_DIR}/preview (2).png", width=Inches(5.0))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 11 — RESULTS & DISCUSSION (Part 2)
# ══════════════════════════════════════════════════════════════════════════════
section_header("6.2 Resume Analysis — Line-by-Line Bullet Feedback")
body(
    "The Sentence Transformer Cosine Similarity engine evaluates each bullet point against "
    "the Gold Standard matrix and generates specific, actionable NLP feedback. Figure 6.2 "
    "shows representative output from the Deep Resume Analysis tab."
)
add_img(f"{IMG_DIR}/preview (3).png", width=Inches(5.5))

doc.add_paragraph()

section_header("6.3 JD Matcher — Cosine Similarity Score")
body(
    "The JD Matcher tab allows users to paste any job description. The system immediately "
    "builds a TF-IDF vector space and computes the Cosine Similarity between the resume "
    "and the JD. Figure 6.3 shows the Jobscan Matcher output for a sample JD."
)
add_img(f"{IMG_DIR}/preview (13).png", width=Inches(5.0))
add_img(f"{IMG_DIR}/preview (14).png", width=Inches(5.0))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 12 — LEARNING OUTCOMES
# ══════════════════════════════════════════════════════════════════════════════
heading("7. LEARNING OUTCOMES", size=14)
divider()
doc.add_paragraph()

section_header("7.1 Technical Skills Gained")
body("This internship provided hands-on experience in the following technical domains:")

learning = [
    ("Information Retrieval & NLP",
     "Implemented TF-IDF from scratch within the scikit-learn framework, "
     "deeply understanding IDF weighting, sublinear TF scaling, and how Cosine Similarity "
     "works in high-dimensional sparse vector spaces."),
    ("Machine Learning — Classification",
     "Trained, evaluated, and serialised a Multinomial Naive Bayes model. Understood the "
     "probabilistic foundations of the Naive Bayes theorem and its practical effectiveness "
     "for text classification tasks."),
    ("Dense Vector Representations (Embeddings)",
     "Worked with the SentenceTransformers library to encode natural language into dense "
     "384-dimensional embeddings using a pre-trained transformer model (all-MiniLM-L6-v2), "
     "and used these for fast approximate nearest-neighbor search via Cosine Similarity."),
    ("Full-Stack Web Development with Python",
     "Designed and built a fully functional, multi-page web application using Streamlit, "
     "with a custom Vanilla CSS glassmorphic design system, session state management, "
     "and dynamic widget rendering."),
    ("PDF Processing & Generation",
     "Used PyMuPDF (fitz) for robust PDF text extraction and FPDF2 / Playwright for "
     "programmatic generation of high-fidelity ATS-ready PDF documents."),
    ("Database Engineering",
     "Designed and queried a SQLite relational database containing 2,800+ normalised "
     "job listings with filters for location, skills, and role."),
]
for title, desc in learning:
    p2 = doc.add_paragraph()
    run_t = p2.add_run(f"{title}: ")
    set_font(run_t, size=12, bold=True)
    run_d = p2.add_run(desc)
    set_font(run_d, size=12)
    p2.paragraph_format.space_after = Pt(5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 13 — CONCLUSION & FUTURE SCOPE
# ══════════════════════════════════════════════════════════════════════════════
heading("8. CONCLUSION", size=14)
divider()
doc.add_paragraph()

body(
    "ResumeIQ successfully demonstrates how a thoughtfully designed ensemble of classical "
    "Machine Learning models and modern NLP techniques can address a real-world problem at "
    "production scale. The Job Role Classifier achieves a 91.69% test accuracy on 25 "
    "industry categories. The TF-IDF + Cosine Similarity job matcher provides semantically "
    "meaningful rankings across a corpus of 2,800+ real job listings. The Sentence Transformer "
    "Cosine Similarity engine delivers granular, actionable feedback on resume bullet quality."
)
body(
    "The platform has been successfully deployed to the Streamlit Community Cloud and is "
    "publicly accessible. The project demonstrates an end-to-end ML workflow — from data "
    "collection, model training, and serialization, to deployment and user interaction — "
    "entirely within a Python ecosystem."
)

doc.add_paragraph()
heading("9. FUTURE SCOPE", size=14)
divider()
doc.add_paragraph()

future = [
    "Replace the Naive Bayes classifier with a fine-tuned BERT-based model for richer contextual understanding of resume text.",
    "Integrate a Retrieval-Augmented Generation (RAG) pipeline to dynamically source job descriptions from live APIs (LinkedIn, Indeed) rather than a static SQLite corpus.",
    "Add a Resume Tailoring Agent that automatically rewrites bullet points to target a specific job description using a fine-tuned instruction-following LLM.",
    "Extend the platform to support DOCX resume uploads in addition to PDF.",
    "Implement a collaborative Leaderboard and peer-review system where users can anonymously compare and rate each other's resumes.",
    "Deploy the backend ML inference pipeline as a separate FastAPI microservice for horizontal scalability.",
]
for f in future:
    bullet(f)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 14 — REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading("10. REFERENCES", size=14)
divider()
doc.add_paragraph()

refs = [
    "[1] Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. "
    "Journal of Machine Learning Research, 12, 2825–2830.",
    "[2] Reimers, N. & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using "
    "Siamese BERT-Networks. Proceedings of EMNLP 2019. arXiv:1908.10084.",
    "[3] Bird, S., Klein, E., & Loper, E. (2009). Natural Language Processing with Python. "
    "O'Reilly Media.",
    "[4] Salton, G. & Buckley, C. (1988). Term-Weighting Approaches in Automatic Text "
    "Retrieval. Information Processing & Management, 24(5), 513–523.",
    "[5] Manning, C. D., Raghavan, P., & Schütze, H. (2008). Introduction to Information "
    "Retrieval. Cambridge University Press.",
    "[6] Streamlit Inc. (2024). Streamlit Documentation. https://docs.streamlit.io",
    "[7] PyMuPDF Documentation. (2024). https://pymupdf.readthedocs.io",
    "[8] Devlin, J. et al. (2019). BERT: Pre-training of Deep Bidirectional Transformers "
    "for Language Understanding. Proceedings of NAACL-HLT 2019.",
    "[9] 611noorsaeed. Updated Resume Dataset. Kaggle. "
    "https://github.com/611noorsaeed/Resume-Screening-App",
    "[10] Groq Inc. (2024). Groq API Documentation. https://console.groq.com/docs",
]
for ref in refs:
    p2 = doc.add_paragraph()
    run2 = p2.add_run(ref)
    set_font(run2, size=11)
    p2.paragraph_format.space_after = Pt(5)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"  Report saved to: {OUT_PATH}")
